import openpyxl
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import os
import shutil

#Establece un color a la palabra Risk dependiendo de la criticidad
def obtener_estilo_por_riesgo(riesgo):
    estilos = {
        'Critical': {'color': RGBColor(255, 0, 0)},  # Rojo fuerte
        'High': {'color': RGBColor(255, 0, 0)},  # Rojo
        'Medium': {'color': RGBColor(255, 165, 0)},  # Naranja
        'Low': {'color': RGBColor(255, 255, 0)},  # Amarillo
        'None': {'color': RGBColor(0, 0, 255)}  # Azul
    }
    return estilos.get(riesgo, {'color': RGBColor(0, 0, 0)})  # Negro por defecto

def agregar_datos_a_documento_existente(ruta_excel, ruta_documento_existente, directorio_destino):
    try:
        print("Cargando el archivo Excel...")
        libro_excel = openpyxl.load_workbook(ruta_excel)
        hoja = libro_excel.active

        ip_columna = 1  # Supongamos que la columna de IPs es la primera (columna A)

        # Obtener IPs únicas para generar los documentos
        ips_unicas = set()
        for row in range(2, hoja.max_row + 1):
            ip = hoja.cell(row=row, column=ip_columna).value
            if ip is not None:
                ips_unicas.add(ip)

        # Copiar el documento existente a un directorio temporal
        shutil.copy(ruta_documento_existente, 'documento_temporal.docx')

        # Para cada IP única, agregar los datos correspondientes al documento existente
        for ip in ips_unicas:
            print(f"Procesando datos para la IP: {ip}")

            # Cargar el documento existente
            doc_existente = Document('documento_temporal.docx')

            # Agregar la IP al documento existente
            p = doc_existente.add_paragraph(f"Datos para la IP: {ip}")
            p.style.font.name = 'Consolas'

            # Agregar datos correspondientes a la IP al documento existente
            for row in range(2, hoja.max_row + 1):
                if hoja.cell(row=row, column=ip_columna).value == ip:
                    for col in range(2, hoja.max_column + 1):
                        nombre_col = hoja.cell(row=1, column=col).value
                        riesgo = hoja.cell(row=row, column=col).value
                        if nombre_col is not None and riesgo is not None:
                            if nombre_col.lower() == 'name':
                                # Agregar un párrafo vacío antes de "Name:"
                                doc_existente.add_paragraph().clear()
                                # Agregar "Name:" en negrita
                                p = doc_existente.add_paragraph()
                                p.add_run('Name: ').bold = True
                                p.add_run(str(riesgo)).font.name = 'Consolas'
                            else:
                                p = doc_existente.add_paragraph()
                                p.add_run(f"{nombre_col}: ").font.color.rgb = obtener_estilo_por_riesgo(riesgo)['color']
                                p.add_run(str(riesgo)).font.name = 'Consolas'
                            # Establecer tamaño de fuente en 9 puntos
                            p.runs[-1].font.size = Pt(9)

            # Guardar el documento con los datos de la IP
            nombre_documento = f"SutituirPorNombre_{ip}.docx" #Nombre de los los documentos que se van a generar.
            ruta_nuevo_documento = os.path.join(directorio_destino, nombre_documento)
            doc_existente.save(ruta_nuevo_documento)
            print(f"Documento para la IP {ip} generado y guardado en: {ruta_nuevo_documento}")

        print("Proceso completado. Documentos Word generados exitosamente.")

    except Exception as e:
        print(f"Error: {str(e)}")
    finally:
        # Eliminar el documento temporal
        os.remove('documento_temporal.docx')

#Estos datos se modifican para indicar el archivo de origen xlsx o csv, el archivo donde se pegara la informacion .docx y la ruta destino donde se guardaran los archivos creados
if __name__ == "__main__":
    ruta_excel = 'Sustituir por ruta' #Ruta del csv o excel que genero la herramienta.
    ruta_documento_existente = 'Sustituir por ruta' #Ruta del documento existente a editar.
    directorio_destino = 'Sustituir por ruta' #Ruta donde se guardaran los archivos generados.
    agregar_datos_a_documento_existente(ruta_excel, ruta_documento_existente, directorio_destino)
